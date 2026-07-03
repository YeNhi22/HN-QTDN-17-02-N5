# -*- coding: utf-8 -*-
"""
Script tạo báo cáo bài tập lớn định dạng Word
Font: Times New Roman, Size: 13
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

def set_font(run, font_name='Times New Roman', font_size=13, bold=False, italic=False):
    """Thiết lập font chữ cho đoạn text"""
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.bold = bold
    run.italic = italic
    # Set font for complex scripts (Vietnamese)
    r = run._element
    rFonts = r.rPr.rFonts
    rFonts.set(qn('w:eastAsia'), font_name)

def add_heading_custom(doc, text, level=1):
    """Thêm heading với font Times New Roman"""
    if level == 1:
        heading = doc.add_heading(text, level=1)
        for run in heading.runs:
            set_font(run, font_size=16, bold=True)
    elif level == 2:
        heading = doc.add_heading(text, level=2)
        for run in heading.runs:
            set_font(run, font_size=14, bold=True)
    elif level == 3:
        heading = doc.add_heading(text, level=3)
        for run in heading.runs:
            set_font(run, font_size=13, bold=True)
    return heading

def add_paragraph_custom(doc, text, bold=False, italic=False, alignment=None):
    """Thêm paragraph với font Times New Roman 13"""
    para = doc.add_paragraph()
    run = para.add_run(text)
    set_font(run, bold=bold, italic=italic)
    if alignment:
        para.alignment = alignment
    return para

def create_cover_page(doc):
    """Tạo trang bìa"""
    # Logo trường (nếu có)
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Tên trường
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('TRƯỜNG ĐẠI HỌC ĐÀ NẴNG\n')
    set_font(run, font_size=14, bold=True)
    run = p.add_run('KHOA CÔNG NGHỆ THÔNG TIN')
    set_font(run, font_size=14, bold=True)
    
    # Thêm khoảng trắng
    for _ in range(3):
        doc.add_paragraph()
    
    # Tiêu đề chính
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('BÁO CÁO BÀI TẬP LỚN')
    set_font(run, font_size=18, bold=True)
    
    # Môn học
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('MÔN: THỰC TẬP CNTT7')
    set_font(run, font_size=16, bold=True)
    
    # Khoảng trắng
    doc.add_paragraph()
    
    # Đề tài
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('ĐỀ TÀI:\nNÂNG CẤP HỆ THỐNG ERP TRÊN ODOO 15\n')
    set_font(run, font_size=15, bold=True)
    run = p.add_run('QUẢN LÝ TÀI SẢN + TÀI CHÍNH/KẾ TOÁN + HRM')
    set_font(run, font_size=15, bold=True)
    
    # Khoảng trắng
    for _ in range(3):
        doc.add_paragraph()
    
    # Thông tin sinh viên
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('SINH VIÊN THỰC HIỆN: NHÓM BTL - FIT-DNU\n')
    set_font(run, font_size=13)
    run = p.add_run('LỚP: K15\n')
    set_font(run, font_size=13)
    run = p.add_run('GIẢNG VIÊN HƯỚNG DẪN: [Tên GVHD]')
    set_font(run, font_size=13)
    
    # Khoảng trắng
    for _ in range(3):
        doc.add_paragraph()
    
    # Thời gian
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Đà Nẵng, tháng 7 năm 2025')
    set_font(run, font_size=13, italic=True)
    
    # Ngắt trang
    doc.add_page_break()

def create_table_of_contents(doc):
    """Tạo mục lục"""
    add_heading_custom(doc, 'MỤC LỤC', level=1)
    add_paragraph_custom(doc, 'CHƯƠNG 1: TỔNG QUAN HỆ THỐNG .................................................. 3', alignment=WD_ALIGN_PARAGRAPH.LEFT)
    add_paragraph_custom(doc, 'CHƯƠNG 2: AUDIT CODE & GAP ANALYSIS .................................... 4', alignment=WD_ALIGN_PARAGRAPH.LEFT)
    add_paragraph_custom(doc, 'CHƯƠNG 3: BUSINESS WORKFLOW ................................................ 15', alignment=WD_ALIGN_PARAGRAPH.LEFT)
    add_paragraph_custom(doc, 'CHƯƠNG 4: IMPLEMENTATION ....................................................... 25', alignment=WD_ALIGN_PARAGRAPH.LEFT)
    add_paragraph_custom(doc, 'CHƯƠNG 5: KẾT LUẬN VÀ HƯỚNG PHÁT TRIỂN ........................... 30', alignment=WD_ALIGN_PARAGRAPH.LEFT)
    add_paragraph_custom(doc, 'PHỤ LỤC: DANH SÁCH FILE CODE ĐÃ SỬA ................................. 31', alignment=WD_ALIGN_PARAGRAPH.LEFT)
    doc.add_page_break()

def create_chapter1(doc):
    """Chương 1: Tổng quan hệ thống"""
    add_heading_custom(doc, 'CHƯƠNG 1: TỔNG QUAN HỆ THỐNG', level=1)
    
    add_heading_custom(doc, '1.1. Giới thiệu đề tài', level=2)
    add_paragraph_custom(doc, 
        'Đề tài "Nâng cấp hệ thống ERP trên Odoo 15" tập trung vào việc tích hợp 3 module chính: '
        'Quản lý Nhân sự (HRM), Quản lý Tài sản và Quản lý Tài chính/Kế toán. Mục tiêu là xây dựng '
        'một hệ thống quản lý doanh nghiệp tích hợp, giảm thiểu nhập liệu trùng lặp và tự động hóa '
        'các quy trình nghiệp vụ.')
    
    add_heading_custom(doc, '1.2. Mục tiêu dự án', level=2)
    add_paragraph_custom(doc, 'Dự án hướng tới các mục tiêu sau:')
    add_paragraph_custom(doc, '• Mức 1 - Tích hợp hệ thống: Đồng bộ dữ liệu giữa các module, loại bỏ nhập liệu thủ công')
    add_paragraph_custom(doc, '• Mức 2 - Tự động hóa quy trình: Tự động tính khấu hao, ghi nhận sổ cái, phê duyệt')
    add_paragraph_custom(doc, '• Áp dụng chuẩn Odoo 15: Sử dụng ORM, computed fields, constraints')
    add_paragraph_custom(doc, '• Đảm bảo bảo mật: Phân quyền theo vai trò, audit trail đầy đủ')
    
    add_heading_custom(doc, '1.3. Công nghệ sử dụng', level=2)
    add_paragraph_custom(doc, '• Nền tảng: Python Odoo 15')
    add_paragraph_custom(doc, '• Database: PostgreSQL')
    add_paragraph_custom(doc, '• Frontend: Odoo Web Framework (JavaScript, XML)')
    add_paragraph_custom(doc, '• Version Control: Git/GitHub')
    
    add_heading_custom(doc, '1.4. Kiến trúc hệ thống', level=2)
    add_paragraph_custom(doc, 
        'Hệ thống được thiết kế theo kiến trúc 3 lớp (3-tier architecture) của Odoo:')
    add_paragraph_custom(doc, '• Presentation Layer: Odoo Web Client (views XML)')
    add_paragraph_custom(doc, '• Business Logic Layer: Python Models với ORM')
    add_paragraph_custom(doc, '• Data Layer: PostgreSQL Database')
    
    doc.add_page_break()
